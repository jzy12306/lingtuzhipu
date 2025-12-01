import io
import re
import tempfile
import os
from typing import Optional
from PyPDF2 import PdfReader
import pandas as pd
import docx
from src.services.ocr_service import ocr_service


def clean_text(text: str) -> str:
    """
    清理文本内容，去除多余的空白字符和特殊字符
    
    Args:
        text: 原始文本
        
    Returns:
        清理后的文本
    """
    # 去除多余的空白字符
    text = re.sub(r'\s+', ' ', text)
    # 去除首尾空白
    text = text.strip()
    # 去除不可见的控制字符（保留换行和制表符）
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t')
    return text


async def process_pdf(file_stream: io.BytesIO) -> str:
    """
    处理PDF文件，提取文本内容，支持OCR识别
    
    Args:
        file_stream: PDF文件流
        
    Returns:
        提取的文本内容
    """
    try:
        import logging
        logger = logging.getLogger(__name__)
        
        logger.info("开始处理PDF文件")
        reader = PdfReader(file_stream)
        text_content = []
        
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            page_text = page.extract_text()
            if page_text:
                text_content.append(clean_text(page_text))
        
        extracted_text = '\n\n'.join(text_content)
        logger.info(f"PDF提取文本长度: {len(extracted_text.strip())}字符")
        
        # 如果提取的文本很少，可能是扫描件，使用OCR识别
        if len(extracted_text.strip()) < 100:
            logger.info(f"PDF提取文本较少（{len(extracted_text.strip())}字符），使用OCR识别")
            # 重置文件流位置
            file_stream.seek(0)
            # 使用OCR服务识别PDF
            ocr_result = await ocr_service.recognize_pdf_stream(file_stream)
            logger.info(f"OCR调用结果: {ocr_result}")
            
            # 解析OCR结果 - 增强版，支持多种返回格式
            ocr_text = []
            
            # 检查OCR结果格式
            if isinstance(ocr_result, dict):
                # 处理有道智云OCR API返回格式
                if 'regions' in ocr_result:
                    # 标准格式，包含regions字段
                    regions = ocr_result.get('regions', [])
                    logger.info(f"OCR识别到的区域数量: {len(regions)}")
                    for region in regions:
                        lines = region.get('lines', [])
                        for line in lines:
                            words = line.get('words', [])
                            for word in words:
                                ocr_text.append(word.get('text', ''))
                elif 'Result' in ocr_result:
                    # 有道智云OCR API可能返回的格式
                    result_text = ocr_result.get('Result', '')
                    logger.info(f"OCR识别到的文本: {result_text}")
                    ocr_text.append(result_text)
                elif 'content' in ocr_result:
                    # 自定义格式，包含content字段
                    content_text = ocr_result.get('content', '')
                    logger.info(f"OCR识别到的内容: {content_text}")
                    ocr_text.append(content_text)
                else:
                    # 其他格式，尝试直接获取文本
                    logger.warning(f"未知的OCR结果格式: {ocr_result.keys()}")
                    # 尝试将整个结果转换为字符串
                    ocr_text.append(str(ocr_result))
            else:
                # 非字典格式，直接转换为字符串
                logger.warning(f"OCR结果不是字典格式: {type(ocr_result)}")
                ocr_text.append(str(ocr_result))
            
            ocr_extracted_text = clean_text(' '.join(ocr_text))
            logger.info(f"OCR识别文本长度: {len(ocr_extracted_text.strip())}字符")
            
            # 如果OCR识别到了文本，使用OCR结果
            if len(ocr_extracted_text.strip()) > 0:
                return ocr_extracted_text
        
        return extracted_text
        
    except Exception as e:
        logger.error(f"PDF处理失败，尝试使用OCR: {str(e)}")
        # 如果正常PDF处理失败，尝试使用OCR
        try:
            # 重置文件流位置
            file_stream.seek(0)
            # 使用OCR服务识别PDF
            ocr_result = await ocr_service.recognize_pdf_stream(file_stream)
            logger.info(f"OCR调用结果: {ocr_result}")
            
            # 解析OCR结果 - 增强版，支持多种返回格式
            ocr_text = []
            
            # 检查OCR结果格式
            if isinstance(ocr_result, dict):
                # 处理有道智云OCR API返回格式
                if 'regions' in ocr_result:
                    # 标准格式，包含regions字段
                    regions = ocr_result.get('regions', [])
                    logger.info(f"OCR识别到的区域数量: {len(regions)}")
                    for region in regions:
                        lines = region.get('lines', [])
                        for line in lines:
                            words = line.get('words', [])
                            for word in words:
                                ocr_text.append(word.get('text', ''))
                elif 'Result' in ocr_result:
                    # 有道智云OCR API可能返回的格式
                    result_text = ocr_result.get('Result', '')
                    logger.info(f"OCR识别到的文本: {result_text}")
                    ocr_text.append(result_text)
                elif 'content' in ocr_result:
                    # 自定义格式，包含content字段
                    content_text = ocr_result.get('content', '')
                    logger.info(f"OCR识别到的内容: {content_text}")
                    ocr_text.append(content_text)
                else:
                    # 其他格式，尝试直接获取文本
                    logger.warning(f"未知的OCR结果格式: {ocr_result.keys()}")
                    # 尝试将整个结果转换为字符串
                    ocr_text.append(str(ocr_result))
            else:
                # 非字典格式，直接转换为字符串
                logger.warning(f"OCR结果不是字典格式: {type(ocr_result)}")
                ocr_text.append(str(ocr_result))
            
            ocr_extracted_text = clean_text(' '.join(ocr_text))
            logger.info(f"OCR识别文本长度: {len(ocr_extracted_text.strip())}字符")
            return ocr_extracted_text
        except Exception as ocr_e:
            logger.error(f"OCR尝试也失败: {str(ocr_e)}")
            # 即使OCR失败，也返回原始提取的文本
            return extracted_text


async def process_text(file_stream: io.BytesIO) -> str:
    """
    处理文本文件，提取文本内容
    
    Args:
        file_stream: 文本文件流
        
    Returns:
        提取的文本内容
    """
    try:
        # 尝试不同的编码
        encodings = ['utf-8', 'gbk', 'latin-1']
        
        for encoding in encodings:
            try:
                file_stream.seek(0)
                text = file_stream.read().decode(encoding)
                return clean_text(text)
            except UnicodeDecodeError:
                continue
        
        raise Exception("无法解码文本文件，请检查文件编码")
        
    except Exception as e:
        raise Exception(f"处理文本文件失败: {str(e)}")


async def process_csv(file_stream: io.BytesIO) -> str:
    """
    处理CSV文件，提取文本内容
    
    Args:
        file_stream: CSV文件流
        
    Returns:
        提取的文本内容
    """
    try:
        file_stream.seek(0)
        # 尝试不同的编码读取CSV
        encodings = ['utf-8', 'gbk', 'latin-1']
        df = None
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_stream)
                break
            except UnicodeDecodeError:
                file_stream.seek(0)
                continue
        
        if df is None:
            raise Exception("无法解码CSV文件，请检查文件编码")
        
        # 将DataFrame转换为文本描述
        text_content = []
        
        # 添加列信息
        text_content.append(f"列名: {', '.join(df.columns.tolist())}")
        text_content.append("")
        
        # 添加数据类型信息
        dtypes_info = []
        for col in df.columns:
            dtypes_info.append(f"{col}: {str(df[col].dtype)}")
        text_content.append(f"数据类型: {'; '.join(dtypes_info)}")
        text_content.append("")
        
        # 添加数据统计信息
        text_content.append(f"总行数: {len(df)}")
        text_content.append("")
        
        # 添加前10行数据作为示例
        text_content.append("数据示例 (前10行):")
        text_content.append(df.head(10).to_string())
        
        # 添加数值列的统计信息
        numeric_cols = df.select_dtypes(include=['number']).columns
        if len(numeric_cols) > 0:
            text_content.append("")
            text_content.append("数值列统计信息:")
            text_content.append(df[numeric_cols].describe().to_string())
        
        return '\n'.join(text_content)
        
    except Exception as e:
        raise Exception(f"处理CSV文件失败: {str(e)}")


async def process_excel(file_stream: io.BytesIO) -> str:
    """
    处理Excel文件，提取文本内容
    
    Args:
        file_stream: Excel文件流
        
    Returns:
        提取的文本内容
    """
    try:
        file_stream.seek(0)
        xl = pd.ExcelFile(file_stream)
        text_content = []
        
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            
            text_content.append(f"\n=== 工作表: {sheet_name} ===")
            text_content.append(f"列名: {', '.join(df.columns.tolist())}")
            text_content.append(f"总行数: {len(df)}")
            
            # 添加前5行数据作为示例
            if len(df) > 0:
                text_content.append("数据示例 (前5行):")
                text_content.append(df.head(5).to_string())
            
            # 添加数值列的统计信息
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_content.append("数值列统计信息:")
                text_content.append(df[numeric_cols].describe().to_string())
        
        return '\n'.join(text_content)
        
    except Exception as e:
        raise Exception(f"处理Excel文件失败: {str(e)}")


async def process_docx(file_stream: io.BytesIO) -> str:
    """
    处理Word文档，提取文本内容
    
    Args:
        file_stream: Word文档流
        
    Returns:
        提取的文本内容
    """
    try:
        doc = docx.Document(file_stream)
        text_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(clean_text(para.text))
        
        # 添加表格信息
        for i, table in enumerate(doc.tables):
            text_content.append(f"\n=== 表格 {i+1} ===")
            
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(clean_text(cell.text))
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n\n'.join(text_content)
        
    except Exception as e:
        raise Exception(f"处理Word文档失败: {str(e)}")


async def process_image(file_stream: io.BytesIO, content_type: str) -> str:
    """
    处理图片文件，使用OCR识别文本内容
    
    Args:
        file_stream: 图片文件流
        content_type: 图片内容类型
        
    Returns:
        OCR识别的文本内容
    """
    try:
        import logging
        logger = logging.getLogger(__name__)
        
        # 重置文件流位置
        file_stream.seek(0)
        
        # 直接使用文件流调用OCR服务
        ocr_result = await ocr_service.recognize_image_stream(file_stream)
        logger.info(f"图片OCR调用结果: {ocr_result}")
        
        # 解析OCR结果 - 增强版，支持多种返回格式
        text_content = []
        
        # 检查OCR结果格式
        if isinstance(ocr_result, dict):
            # 处理有道智云OCR API返回格式
            if 'regions' in ocr_result:
                # 标准格式，包含regions字段
                regions = ocr_result.get('regions', [])
                logger.info(f"图片OCR识别到的区域数量: {len(regions)}")
                for region in regions:
                    lines = region.get('lines', [])
                    for line in lines:
                        words = line.get('words', [])
                        for word in words:
                            text_content.append(word.get('text', ''))
            elif 'Result' in ocr_result:
                # 有道智云OCR API可能返回的格式
                result_text = ocr_result.get('Result', '')
                logger.info(f"图片OCR识别到的文本: {result_text}")
                text_content.append(result_text)
            elif 'content' in ocr_result:
                # 自定义格式，包含content字段
                content_text = ocr_result.get('content', '')
                logger.info(f"图片OCR识别到的内容: {content_text}")
                text_content.append(content_text)
            else:
                # 其他格式，尝试直接获取文本
                logger.warning(f"未知的图片OCR结果格式: {ocr_result.keys()}")
                # 尝试将整个结果转换为字符串
                text_content.append(str(ocr_result))
        else:
            # 非字典格式，直接转换为字符串
            logger.warning(f"图片OCR结果不是字典格式: {type(ocr_result)}")
            text_content.append(str(ocr_result))
        
        recognized_text = clean_text(' '.join(text_content))
        logger.info(f"图片OCR识别文本长度: {len(recognized_text.strip())}字符")
        return recognized_text
            
    except Exception as e:
        logger.error(f"处理图片文件失败: {str(e)}")
        # 即使处理失败，也返回空字符串，避免整个流程中断
        return ""


async def process_uploaded_file(file) -> tuple[str, str]:
    """
    根据文件类型处理上传的文件，提取文本内容
    
    Args:
        file: UploadFile 对象
        
    Returns:
        tuple: (提取的文本内容, 文件类型)
    """
    # 获取文件流和内容类型
    file_stream = io.BytesIO(await file.read())
    content_type = file.content_type
    
    # 根据文件名获取文件类型
    file_extension = ''
    if file.filename:
        file_extension = file.filename.split('.')[-1].lower()
    
    # 根据文件扩展名和内容类型选择处理器
    if file_extension == 'pdf' or content_type == 'application/pdf':
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_pdf(file_stream)
        return text_content, 'pdf'
    elif file_extension == 'txt' or content_type == 'text/plain':
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_text(file_stream)
        return text_content, 'txt'
    elif file_extension == 'md' or content_type == 'text/markdown':
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_text(file_stream)
        return text_content, 'txt'  # 将markdown类型统一为txt
    elif file_extension == 'csv' or content_type == 'text/csv':
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_csv(file_stream)
        return text_content, 'txt'  # CSV文件作为txt处理
    elif file_extension in ['xls', 'xlsx'] or content_type in [
        'application/vnd.ms-excel',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    ]:
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_excel(file_stream)
        return text_content, 'txt'  # Excel文件作为txt处理
    elif file_extension == 'docx' or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_docx(file_stream)
        return text_content, 'docx'
    elif file_extension in ['jpg', 'jpeg', 'png'] or content_type in ['image/jpeg', 'image/png', 'image/jpg']:
        # 图片文件，使用OCR识别
        # 在调用处理器函数之前重置文件流位置
        file_stream.seek(0)
        text_content = await process_image(file_stream, content_type)
        return text_content, 'txt'  # 图片OCR结果作为txt处理
    else:
        # 对于无法识别的文件类型，尝试作为文本处理
        try:
            # 在调用处理器函数之前重置文件流位置
            file_stream.seek(0)
            text_content = await process_text(file_stream)
            return text_content, 'txt'
        except Exception as e:
            # 如果作为文本处理失败，返回空字符串和txt类型
            return "", 'txt'


async def get_file_metadata(file_stream: io.BytesIO, content_type: str) -> dict:
    """
    获取文件的元数据信息
    
    Args:
        file_stream: 文件流
        content_type: 文件内容类型
        
    Returns:
        元数据字典
    """
    metadata = {
        "content_type": content_type,
        "file_size": file_stream.getbuffer().nbytes,
    }
    
    # 根据文件类型获取额外的元数据
    if content_type == 'application/pdf':
        try:
            file_stream.seek(0)
            reader = PdfReader(file_stream)
            metadata["page_count"] = len(reader.pages)
            # 尝试获取PDF的元数据
            pdf_metadata = reader.metadata
            if pdf_metadata:
                for key, value in pdf_metadata.items():
                    metadata[f"pdf_{key}"] = value
        except Exception:
            pass
    
    return metadata